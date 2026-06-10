#!/usr/bin/env python3
"""Create a Word shot-list table from shot_list.json."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


COLUMNS = [
    ("Shot", "shot_id"),
    ("Duration", "duration_sec"),
    ("Ref Beat", "reference_beat_id"),
    ("Shot Size", "shot_size"),
    ("Visual", "visual"),
    ("Action", "action"),
    ("Camera", "camera"),
    ("Product Role", "product_role"),
    ("On-screen Text", "on_screen_text"),
    ("Source Refs", "source_refs"),
    ("AI Notes", "ai_video_notes"),
    ("Risk Flags", "risk_flags"),
]


def stringify(value: object) -> str:
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    return "" if value is None else str(value)


def create_docx(input_json: Path, output_docx: Path) -> None:
    data = json.loads(input_json.read_text(encoding="utf-8"))
    shots = data.get("shots", [])
    if not isinstance(shots, list) or not shots:
        raise ValueError("shot_list.json must contain a non-empty shots array")

    document = Document()
    document.add_heading(data.get("project_title", "Shot List"), level=1)
    duration = data.get("target_duration_sec")
    if duration:
        document.add_paragraph(f"Target duration: {duration}s")

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for index, (label, _) in enumerate(COLUMNS):
        header[index].text = label

    for shot in shots:
        row = table.add_row().cells
        for index, (_, key) in enumerate(COLUMNS):
            row[index].text = stringify(shot.get(key))

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: create_shot_list_docx.py <shot_list.json> <shot_list.docx>")
        return 2
    create_docx(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"Created {sys.argv[2]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
